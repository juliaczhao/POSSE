"""Per-program Word report.

Combines the program JSON, the LLM summary CSVs, and the report plots into a
single .docx under <working_dir>/reports/.
"""

import logging
import json
import csv
import os
import anndata
from typing import Dict

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


TABLE_HEADER_BLUE_HEX = "D9EAF7"


def _set_cell_shading(cell, hex_color: str):
    """Set background shading on a docx table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


class ProgramReportGenerator:
    """Generate beautifully formatted Word documents for gene programs."""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Create custom styles for the document."""
        styles = self.doc.styles
        
        if 'Main Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Main Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
        if 'Subtitle' not in [s.name for s in styles]:
            subtitle_style = styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Calibri'
            subtitle_font.size = Pt(16)
            subtitle_font.bold = True
            subtitle_font.color.rgb = RGBColor(51, 51, 51)
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(8)
            
        if 'Program Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Program Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(18)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 102, 204)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
        if 'Section Header' not in [s.name for s in styles]:
            header_style = styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Calibri'
            header_font.size = Pt(14)
            header_font.bold = True
            header_font.color.rgb = RGBColor(51, 51, 51)
            header_style.paragraph_format.space_before = Pt(12)
            header_style.paragraph_format.space_after = Pt(6)
            
        if 'Description Text' not in [s.name for s in styles]:
            desc_style = styles.add_style('Description Text', WD_STYLE_TYPE.PARAGRAPH)
            desc_font = desc_style.font
            desc_font.name = 'Calibri'
            desc_font.size = Pt(11)
            desc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            desc_style.paragraph_format.line_spacing = 1.15
            desc_style.paragraph_format.space_after = Pt(12)
            
        if 'Gene List' not in [s.name for s in styles]:
            gene_style = styles.add_style('Gene List', WD_STYLE_TYPE.PARAGRAPH)
            gene_font = gene_style.font
            gene_font.name = 'Consolas'
            gene_font.size = Pt(9)
            gene_font.color.rgb = RGBColor(102, 102, 102)
            gene_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            gene_style.paragraph_format.line_spacing = 1.1
            gene_style.paragraph_format.space_after = Pt(6)
            
        if 'Summary Text' not in [s.name for s in styles]:
            summary_style = styles.add_style('Summary Text', WD_STYLE_TYPE.PARAGRAPH)
            summary_font = summary_style.font
            summary_font.name = 'Calibri'
            summary_font.size = Pt(11)
            summary_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            summary_style.paragraph_format.line_spacing = 1.2
            summary_style.paragraph_format.space_after = Pt(8)

    def load_programs_data(self, working_dir: str) -> Dict:
        """Load programs data from JSON file."""
        programs_path = os.path.join(working_dir, 'clique_based_programs', 'programs_after_grow_merge.json')
        if not os.path.exists(programs_path):
            raise FileNotFoundError(f"Programs file not found: {programs_path}")
        
        with open(programs_path, 'r') as f:
            programs_data = json.load(f)
        
        logger.info(f"Loaded {len(programs_data)} programs from {programs_path}")
        return programs_data

    def _get_program_field(self, row: Dict, *candidates) -> str:
        for cand in candidates:
            for key in row.keys():
                if key.strip().lower() == cand.lower():
                    val = row[key]
                    return val.strip() if isinstance(val, str) else val
        return ""

    def load_program_summaries(self, working_dir: str) -> Dict:
        """Load program summaries from CSV file."""
        summaries_path = os.path.join(working_dir, 'llm_summaries', 'program_summaries_with_genes.csv')
        if not os.path.exists(summaries_path):
            logger.warning(f"Summaries file not found: {summaries_path}")
            return {}
        
        summaries = {}
        with open(summaries_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f, delimiter='$')
            for row in reader:
                program_num = None
                for key in row.keys():
                    if 'program' in key.lower():
                        program_num = row[key].strip()
                        break
                
                if program_num:
                    if program_num.startswith('Program '):
                        program_num = program_num.replace('Program ', '')
                    summaries[program_num] = row
        
        logger.info(f"Loaded {len(summaries)} program summaries from {summaries_path}")
        return summaries

    def get_dataset_info(self, working_dir: str) -> Dict:
        """Get dataset information from AnnData files."""
        anndata_dir = os.path.join(working_dir, 'anndata_files')
        if os.path.exists(anndata_dir):
            raw_files = [f for f in os.listdir(anndata_dir) if f.startswith('raw') and f.endswith('.h5ad')]
            
            if raw_files:
                adata_path = os.path.join(anndata_dir, raw_files[0])
                try:
                    adata = anndata.read_h5ad(adata_path)
                    
                    total_cells = adata.n_obs
                    if len(raw_files) > 1:
                        total_cells = 0
                        for raw_file in raw_files:
                            chunk_path = os.path.join(anndata_dir, raw_file)
                            chunk_adata = anndata.read_h5ad(chunk_path)
                            total_cells += chunk_adata.n_obs
                    
                    return {
                        'n_cells': total_cells,
                        'n_genes': adata.n_vars,
                        'cell_types': list(adata.obs['cell_type'].unique()) if 'cell_type' in adata.obs.columns else []
                    }
                except Exception as e:
                    logger.warning(f"Could not read AnnData file {adata_path}: {e}")
        
        gene_names_dir = os.path.join(working_dir, 'gene_names')
        if os.path.exists(gene_names_dir):
            gene_names_path = os.path.join(gene_names_dir, 'gene_names.npy')
            if os.path.exists(gene_names_path):
                import numpy as np
                gene_names = np.load(gene_names_path, allow_pickle=True)
                return {
                    'n_cells': 'Unknown',
                    'n_genes': len(gene_names),
                    'cell_types': []
                }
        
        return {
            'n_cells': 'Unknown',
            'n_genes': 'Unknown',
            'cell_types': []
        }

    def add_title_page(self, working_dir: str, programs_data: Dict):
        """Add title page with dataset information."""
        self.doc.add_paragraph('Gene Program Analysis Report', style='Main Title')
        
        dataset_info = self.get_dataset_info(working_dir)
        
        self.doc.add_paragraph('Dataset Information', style='Subtitle')
        
        cell_types = dataset_info.get('cell_types', [])
        if not cell_types:
            working_dir_name = os.path.basename(working_dir)
            if working_dir_name != 'all':
                cell_types = [working_dir_name]
        
        cell_type_text = ', '.join(str(cell_type) for cell_type in cell_types) if cell_types else 'All cell types'
        
        details = [
            f"Cell Type(s): {cell_type_text}",
            f"Number of Cells: {dataset_info.get('n_cells', 'Unknown')}",
            f"Number of Genes: {dataset_info.get('n_genes', 'Unknown')}"
        ]
        
        for detail in details:
            p = self.doc.add_paragraph(detail, style='Summary Text')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph('Program Summary', style='Subtitle')
        
        n_programs = len(programs_data)
        all_genes = set()
        program_sizes = []
        
        for program_info in programs_data.values():
            program_genes = program_info.get('program', [])
            program_sizes.append(len(program_genes))
            all_genes.update(program_genes)
        
        n_unique_genes = len(all_genes)
        
        stats = [
            f"Number of Programs Found: {n_programs}",
            f"Number of Unique Genes Covered: {n_unique_genes}",
            f"Program Size Distribution: {sorted(program_sizes, reverse=True)}"
        ]
        
        for stat in stats:
            p = self.doc.add_paragraph(stat, style='Summary Text')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()

    def add_programs_summary_table(self, programs_data: Dict, summaries_data: Dict):
        """Insert a compact 4-column summary table (program #, title, # genes, # inconsistent)."""
        self.doc.add_paragraph('Programs Overview', style='Section Header')

        sorted_programs = sorted(
            programs_data.items(),
            key=lambda x: int(x[0]) if str(x[0]).isdigit() else x[0]
        )

        headers = ["Program", "Title", "# Genes", "# Inconsistent"]
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        col_widths = [Cm(2.4), Cm(8.7), Cm(2.0), Cm(3.0)]

        hdr_cells = table.rows[0].cells
        for i, name in enumerate(headers):
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            _set_cell_shading(hdr_cells[i], TABLE_HEADER_BLUE_HEX)
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hdr_cells[i].width = col_widths[i]

        total_genes = 0
        total_inconsistent = 0

        for program_num, program_info in sorted_programs:
            program_genes = program_info.get('program', [])
            n_genes = len(program_genes)

            row = summaries_data.get(str(program_num), {})
            title = self._get_program_field(row, 'Title') or ''
            n_inc_raw = self._get_program_field(row, 'Number of Inconsistent Genes', 'Num Inconsistent Genes')
            try:
                n_inconsistent = int(n_inc_raw) if n_inc_raw != '' else 0
            except ValueError:
                n_inconsistent = 0

            total_genes += n_genes
            total_inconsistent += n_inconsistent

            cells = table.add_row().cells
            values = [str(program_num), title, str(n_genes), str(n_inconsistent)]
            for i, v in enumerate(values):
                cells[i].text = ""
                p = cells[i].paragraphs[0]
                run = p.add_run(v)
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
                cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[i].width = col_widths[i]

        total_cells = table.add_row().cells
        total_values = ["TOTAL", "", str(total_genes), str(total_inconsistent)]
        for i, v in enumerate(total_values):
            total_cells[i].text = ""
            p = total_cells[i].paragraphs[0]
            run = p.add_run(v)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            total_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            total_cells[i].width = col_widths[i]

        self.doc.add_page_break()

    def add_global_plots_section(self, working_dir: str):
        """Add global plots section."""
        self.doc.add_paragraph('Global Analysis Plots', style='Section Header')
        
        global_plots_dir = os.path.join(working_dir, 'report_plots', 'global_plots')
        if not os.path.exists(global_plots_dir):
            logger.warning(f"Global plots directory not found: {global_plots_dir}")
            return
        
        global_plots = [
            'program_correlation_heatmap.png',
            'global_umap.png',
            'global_umap_leiden.png',
            'cell_type_counts_bar.png',
            'umap_by_cell_type.png',
            'all_programs_violin_plot.png'
        ]
        
        for plot_name in global_plots:
            plot_path = os.path.join(global_plots_dir, plot_name)
            if os.path.exists(plot_path):
                title = plot_name.replace('_', ' ').replace('.png', '').title()
                self.doc.add_paragraph(title, style='Section Header')
                
                self.doc.add_picture(plot_path, width=Inches(6))
                
                self.doc.add_paragraph()
            else:
                logger.warning(f"Global plot not found: {plot_path}")
        
        supp_plots = [
            ('supp_material/program_finding_plots/maximal_cliques_in_programs_final.png', 'Maximal Cliques In Programs Final'),
            ('supp_material/plots/gene_param_distributions.png', 'Gene Parameter Distributions')
        ]
        
        for plot_rel_path, plot_title in supp_plots:
            plot_path = os.path.join(working_dir, plot_rel_path)
            if os.path.exists(plot_path):
                self.doc.add_paragraph(plot_title, style='Section Header')
                
                self.doc.add_picture(plot_path, width=Inches(6))
                
                self.doc.add_paragraph()
            else:
                logger.warning(f"Supplementary plot not found: {plot_path}")
        
        self.doc.add_page_break()

    def add_program_pages(self, working_dir: str, programs_data: Dict, summaries_data: Dict):
        """Add individual program pages."""
        self.doc.add_paragraph('Individual Program Analysis', style='Section Header')
        self.doc.add_page_break()
        
        sorted_programs = sorted(programs_data.items(), key=lambda x: int(x[0]) if x[0].isdigit() else x[0])
        
        for program_num, program_info in sorted_programs:
            self.add_single_program_page(working_dir, program_num, program_info, summaries_data)

    def add_single_program_page(self, working_dir: str, program_num: str, program_info: Dict, summaries_data: Dict):
        """Add a single program page with all required elements."""
        program_genes = program_info.get('program', [])
        llm_title = ""
        if program_num in summaries_data:
            llm_title = self._get_program_field(summaries_data[program_num], 'Title')
        if llm_title:
            title = f'Program {program_num}: {llm_title} ({len(program_genes)} genes)'
        else:
            title = f'Program {program_num} ({len(program_genes)} genes)'
        self.doc.add_paragraph(title, style='Program Title')
        
        if program_num in summaries_data:
            summary = None
            for key in summaries_data[program_num].keys():
                if 'summary' in key.lower():
                    summary = summaries_data[program_num][key]
                    break
            
            if summary and summary.strip():
                self.doc.add_paragraph('Program Description:', style='Section Header')
                self.doc.add_paragraph(summary, style='Description Text')
            else:
                self.doc.add_paragraph('Program Description: No description available.', style='Description Text')
        else:
            self.doc.add_paragraph('Program Description: No description available.', style='Description Text')
        
        violin_path = os.path.join(working_dir, 'report_plots', 'program_plots', 'program_activity_violins', f'program_{program_num}_activity_violin.png')
        if os.path.exists(violin_path):
            self.doc.add_paragraph('Program Activity Across Clusters:', style='Section Header')
            self.doc.add_picture(violin_path, width=Inches(6))
            self.doc.add_paragraph()
        else:
            logger.warning(f"Violin plot not found: {violin_path}")
        
        self.doc.add_paragraph('UMAP Visualizations:', style='Section Header')
        
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        table.autofit = True
        
        global_umap_path = os.path.join(working_dir, 'report_plots', 'global_plots', 'global_umap_leiden.png')
        if os.path.exists(global_umap_path):
            cell = table.cell(0, 0)
            cell.text = 'Global UMAP (Leiden)'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].style = 'Section Header'
            
            cell = table.cell(1, 0)
            run = cell.paragraphs[0].add_run()
            run.add_picture(global_umap_path, width=Inches(2.2))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            logger.warning(f"Global UMAP plot not found: {global_umap_path}")
        
        program_umap_path = os.path.join(working_dir, 'report_plots', 'program_plots', 'program_activity_umaps', f'umap_program_activity_{program_num}.png')
        if os.path.exists(program_umap_path):
            cell = table.cell(0, 1)
            cell.text = f'Program {program_num} Activity UMAP'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].style = 'Section Header'
            
            cell = table.cell(1, 1)
            run = cell.paragraphs[0].add_run()
            run.add_picture(program_umap_path, width=Inches(2.2))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            logger.warning(f"Program UMAP plot not found: {program_umap_path}")
        
        self.doc.add_paragraph()
        
        self.doc.add_paragraph('Program Genes (First 100):', style='Section Header')
        genes_to_show = program_genes[:100]
        genes_text = ', '.join(genes_to_show)
        if len(program_genes) > 100:
            genes_text += f' ... and {len(program_genes) - 100} more genes'
        
        self.doc.add_paragraph(genes_text, style='Gene List')
        
        self.doc.add_page_break()

    def generate_report(self, working_dir: str):
        """Generate the complete program report."""
        logger.info(f"Generating report for working directory: {working_dir}")
        
        programs_data = self.load_programs_data(working_dir)
        summaries_data = self.load_program_summaries(working_dir)
        
        self.add_title_page(working_dir, programs_data)
        self.add_programs_summary_table(programs_data, summaries_data)
        self.add_global_plots_section(working_dir)
        self.add_program_pages(working_dir, programs_data, summaries_data)
        
        output_path = os.path.join(working_dir, 'reports', 'program_analysis_report.docx')
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        
        logger.info(f"Report saved to: {output_path}")


def generate_report(working_dir: str):
    """Main function to generate the program report."""
    generator = ProgramReportGenerator()
    generator.generate_report(working_dir)


